"""Generate pitch script as .docx — Final version matching all 8 required elements"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)

# ─── Title Page ───
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('LY')
r.font.size = Pt(60)
r.font.bold = True
r.font.color.rgb = RGBColor(108, 92, 231)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Smart loyalty for independent businesses')
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(99, 110, 114)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Final Pitch Script')
r.font.size = Pt(14)
r.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('VALI Bootcamp - ESMT Berlin - March 2026')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(99, 110, 114)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Alexander Berner  |  Georgios Zelenitsas  |  Athanasios Manios  |  Romain Bouxirot')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(99, 110, 114)

doc.add_page_break()

# ─── Overview ───
p = doc.add_paragraph()
r = p.add_run('PITCH OVERVIEW')
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = RGBColor(108, 92, 231)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('13 slides  |  ~3:40 total  |  20 seconds buffer')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(99, 110, 114)

doc.add_paragraph()

overview = [
    ('1', 'Title', 'HOOK', '5s'),
    ('2', 'Story Hook', 'HOOK', '30s'),
    ('3', 'The Problem', 'PROBLEM', '25s'),
    ('4', 'Why Solutions Fail', 'PROBLEM', '20s'),
    ('5', 'Our Solution', 'SOLUTION', '15s'),
    ('6', 'Our ICP', 'ICP', '15s'),
    ('7', 'How It Works', 'MVP', '25s'),
    ('8', 'What We Built', 'MVP', '15s'),
    ('9', 'Validation + Pricing', 'MARKET + GTM', '25s'),
    ('10', 'Market Size', 'MARKET', '10s'),
    ('11', 'Go-To-Market', 'GTM', '15s'),
    ('12', 'What We Need', 'ASK', '10s'),
    ('13', 'Closing + Team', 'TEAM', '10s'),
]

table = doc.add_table(rows=len(overview)+1, cols=4)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Slide', 'Title', 'Element', 'Time']):
    table.rows[0].cells[i].text = h
for i, (num, title, elem, time) in enumerate(overview):
    table.rows[i+1].cells[0].text = num
    table.rows[i+1].cells[1].text = title
    table.rows[i+1].cells[2].text = elem
    table.rows[i+1].cells[3].text = time

doc.add_page_break()

# ─── Slides ───
slides = [
    ('1', 'TITLE', '5 seconds',
     'HOOK',
     [
        '[Slide: LY logo on cafe background, tagline]',
        '',
        '"Good morning. We are Team LY."',
     ]),
    ('2', 'THE STORY', '30 seconds',
     'HOOK',
     [
        '[Slide: Dark cinematic background, barista image]',
        '[Slow pace. Eye contact. Let the silence work.]',
        '',
        '"Monday morning. 7:15. Francesca opens her coffee shop in Kreuzberg."',
        '',
        '"Marco should be here by now. Double espresso, no sugar. Every Monday for three years."',
        '',
        '"But Marco didn\'t come last Monday. Or the one before."',
        '',
        '[PAUSE - 2 seconds. Look at audience.]',
        '',
        '"Francesca remembers his birthday. His dog\'s name. The day he switched from cappuccino to espresso."',
        '',
        '"But she had no way of knowing he was leaving."',
        '"No alert. No signal. No second chance."',
     ]),
    ('3', 'THE PROBLEM', '25 seconds',
     'PROBLEM',
     [
        '[Slide: 4 statistics - 10M+, 78%, 5-7x, EUR375K]',
        '',
        '"Now multiply Francesca by 10 million. That\'s how many independent food and beverage businesses across Europe face this exact problem - every single week."',
        '',
        '"78 percent of restaurant guests never come back after their first visit."',
        '',
        '"Each restaurant loses roughly 375,000 euros a year to silent churn. And acquiring a new customer costs 5 to 7 times more than keeping one."',
        '',
        '"These businesses are bleeding revenue - and they don\'t even know it."',
     ]),
    ('4', 'WHY SOLUTIONS FAIL', '20 seconds',
     'PROBLEM',
     [
        '[Slide: Two boxes - stamp cards vs enterprise CRMs]',
        '',
        '"The tools today fall into two buckets."',
        '',
        '"Stamp cards. Buy ten, get one free. No intelligence. They can\'t tell you who\'s leaving."',
        '',
        '"Enterprise CRMs. HubSpot. Salesforce. Thousands per month. Not built for someone who is the barista, the accountant, and the manager - all at once."',
        '',
        '"The gap? No affordable, AI-powered retention tool for independent European F&B."',
        '[Pause]',
        '"Zero."',
     ]),
    ('5', 'OUR SOLUTION', '15 seconds',
     'SOLUTION',
     [
        '[Slide: "Introducing" then LY logo then "Your retention co-pilot"]',
        '',
        '"LY is a retention co-pilot for independent businesses."',
        '',
        '"We don\'t track loyalty points. We watch the patterns you can\'t see. We tell you who\'s leaving, when, and exactly what to do about it."',
        '',
        '"Think of it as a smart marketing assistant - for 29 euros a month."',
     ]),
    ('6', 'OUR ICP', '15 seconds',
     'ICP',
     [
        '[Slide: ICP card with Who/Size/Where/Today/Trigger]',
        '',
        '"Our customer is Francesca. Owner-operator of an independent cafe, restaurant, or bakery. 100 to 500 weekly customers. 1 to 15 employees. Urban European markets."',
        '',
        '"She has a POS but no CRM, no loyalty program, and no data on who\'s leaving. She notices regulars are gone - but only when it\'s too late."',
     ]),
    ('7', 'HOW IT WORKS', '25 seconds',
     'MVP',
     [
        '[Slide: 4 steps with numbers]',
        '',
        '"Here\'s how simple it is."',
        '',
        '"Step 1: Print a QR code. Put it at the counter. Done."',
        '',
        '"Step 2: Customers scan it. Name and phone number. They get their personal loyalty QR code. No app download. 15 seconds."',
        '',
        '"Step 3: Every visit, your employee scans their QR. Visits tracked automatically."',
        '',
        '"Step 4: LY watches the patterns. When Marco hasn\'t been back in 12 days - and he usually comes every 3 - you get an alert. With a suggested action. One WhatsApp message later, Marco is back."',
     ]),
    ('8', 'WHAT WE BUILT', '15 seconds',
     'MVP',
     [
        '[Slide: 6 feature cards + lyloyal.com]',
        '',
        '"We didn\'t just make a pitch deck. We built the product. It\'s live."',
        '',
        '"AI churn detection. Loyalty tiers from Bronze to Platinum. One-click WhatsApp messaging. Custom reward rules. A public shop map. And an AI chatbot."',
        '',
        '"You can try it right now at lyloyal.com."',
        '[Point to the URL on screen]',
     ]),
    ('9', 'VALIDATION + BUSINESS MODEL', '25 seconds',
     'MARKET + GTM',
     [
        '[Slide: Split - interview quotes left, pricing right]',
        '',
        '"We went out and talked to real shop owners in Berlin."',
        '',
        '"How do you know when a regular stops coming?"',
        '[Pause] "They don\'t. They just disappear."',
        '',
        '"What digital tools do you use for retention?"',
        '"None. They can\'t afford them."',
        '',
        '"If we could tell you who\'s about to leave?"',
        '"I\'d call them immediately."',
        '',
        '"Our model: freemium SaaS. Start free with 50 customers. Growth at 29 euros a month for AI churn detection and WhatsApp. Pro at 79 for multi-location."',
        '',
        '"The math is simple: recovering just 2 lost regulars per month pays for the entire subscription."',
     ]),
    ('10', 'MARKET SIZE', '10 seconds',
     'MARKET',
     [
        '[Slide: TAM/SAM/SOM - $6.6B / 5.5M / 275K]',
        '',
        '"The loyalty software market is 6.6 billion dollars, growing at 17% annually. 5.5 million independent F&B businesses in Europe. Our near-term target: 275,000 at 5% penetration. Starting with Berlin."',
     ]),
    ('11', 'GO-TO-MARKET', '15 seconds',
     'GTM',
     [
        '[Slide: 4 GTM steps]',
        '',
        '"Our go-to-market is Berlin first. Door-to-door in Kreuzberg and Neukolln. Walk into 50 cafes with a tablet. Set them up in 2 minutes. Free tier means zero risk."',
        '',
        '"Then pop-up loyalty events. Partner with cafes for LY Day - free coffee for customers who sign up. Viral loop."',
        '',
        '"The shop map creates a network effect. More shops, more customers, more shops. And once owners see churn alerts work - upgrading to Growth is a no-brainer."',
     ]),
    ('12', 'WHAT WE NEED', '10 seconds',
     'ASK',
     [
        '[Slide: 4 icons - Pilots, Mentorship, Market Access, EUR50K]',
        '',
        '"What we need: 10 pilot cafes in Berlin. Mentorship from F&B and SaaS experts through INNOVA EUROPE. Introductions to POS providers like SumUp. And 50,000 euros in pre-seed to hire one engineer and reach 100 shops by Q4."',
     ]),
    ('13', 'CLOSING', '10 seconds',
     'TEAM',
     [
        '[Slide: LY logo, closing line, team names, lyloyal.com]',
        '',
        '[Slow down. Take a breath. Look at the audience.]',
        '',
        '"Because Francesca already has the heart for loyalty."',
        '',
        '[2-second pause]',
        '',
        '"LY just gives her the eyes."',
        '',
        '[Pause. Nod.]',
        '',
        '"Thank you."',
        '',
        '[Wait for applause, then:]',
        '"We are Alexander, Georgios, Athanasios, and Romain. The product is live at lyloyal.com - we\'d love your feedback."',
     ]),
]

for num, title, duration, element, lines in slides:
    # Slide header
    p = doc.add_paragraph()
    r = p.add_run(f'SLIDE {num}')
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(108, 92, 231)
    r = p.add_run(f'  {title}')
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(45, 52, 54)

    # Element badge + duration
    p = doc.add_paragraph()
    r = p.add_run(f'{element}')
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 210, 211)
    r = p.add_run(f'  |  {duration}')
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph()

    for line in lines:
        if not line:
            doc.add_paragraph()
        elif line.startswith('['):
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.italic = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(150, 150, 150)
        else:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(13)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run('_' * 80)
    r.font.color.rgb = RGBColor(220, 220, 220)
    doc.add_paragraph()

doc.add_page_break()

# ─── Delivery Guide ───
p = doc.add_paragraph()
r = p.add_run('DELIVERY GUIDE')
r.font.size = Pt(20)
r.font.bold = True
r.font.color.rgb = RGBColor(108, 92, 231)

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('The 5 Moments That Make or Break This Pitch')
r.font.size = Pt(14)
r.font.bold = True

doc.add_paragraph()

moments = [
    ('The Marco Story (Slide 2)',
     'This is your ENTIRE pitch in miniature. If the audience feels something here, you win. Deliver it slowly. Make eye contact. Let the pauses breathe. The line "she had no way of knowing" should land like a punch. Do NOT rush this.'),
    ('The "Zero" (Slide 4)',
     'After listing the gap, pause. Then say "Zero." alone. One word. Let it hang. This is where the audience realizes the opportunity is real.'),
    ('The Live Product (Slide 8)',
     'This is your unfair advantage. Most bootcamp teams have a slide deck. You have a working product. When you say "try it right now at lyloyal.com" - own it. This is the moment that separates you.'),
    ('The Customer Quotes (Slide 9)',
     'Deliver the quotes like you\'re reading a conversation. Not a slide. "How do you know when regulars leave?" [pause] "They don\'t. They just disappear." Make it feel like the audience is in the room with you and that cafe owner.'),
    ('The Closing Line (Slide 13)',
     '"Francesca already has the heart for loyalty. LY just gives her the eyes." This line is designed to be remembered. Say it slowly. With conviction. Do not add anything after "the eyes" except "Thank you." Less is more.'),
]

for title, body in moments:
    p = doc.add_paragraph()
    r = p.add_run(f'{title}')
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(108, 92, 231)

    p = doc.add_paragraph()
    r = p.add_run(body)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Pre-Pitch Checklist')
r.font.size = Pt(14)
r.font.bold = True

checklist = [
    'lyloyal.com open on a phone (for Q&A demo)',
    'Practice the Marco story 3 times out loud',
    'Time the full pitch - aim for 3:40',
    'Decide who presents which slides if splitting',
    'Test the pitch.html presentation on the venue projector',
    'Have a backup: pitch.html works offline (no internet needed)',
]
for item in checklist:
    doc.add_paragraph(item, style='List Bullet')

doc.save('LY_Pitch_Script.docx')
print('Created LY_Pitch_Script.docx')
